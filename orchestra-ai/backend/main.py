import re
from fastapi import FastAPI, Request, HTTPException, UploadFile, File, WebSocket, WebSocketDisconnect
from fastapi.responses import JSONResponse

from pydantic import BaseModel
from agents import graph_app
import uvicorn
import time
import asyncio
import logging
import json
import os
from fastapi.middleware.cors import CORSMiddleware
from groq import Groq
from dotenv import load_dotenv
from rate_limiter import rate_limiter

load_dotenv()

# --- Configure Background Logging ---
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("orchestra_server.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("orchestra")

app = FastAPI(title="Orchestra AI API")
groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

from fastapi.staticfiles import StaticFiles

# Enable CORS for our Frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Dynamic static file serving with self-healing local hashed image cache
from fastapi.responses import FileResponse
import urllib.parse
import httpx
import random
import hashlib
import os
from fastapi import Response

image_lock = asyncio.Lock()

@app.get("/static/images/{filename}")
async def get_static_image(filename: str, legacy_url: str | None = None):
    from database import db_manager
    
    # 1. Map to safe hashed filename path to prevent illegal characters or path long on Windows
    file_hash = hashlib.sha256(filename.encode('utf-8')).hexdigest()
    # Support both png and jpg (standardized to .jpg as fallback)
    local_path = os.path.join("static", "images", f"{file_hash}.jpg")
    local_path_png = os.path.join("static", "images", f"{file_hash}.png")
    
    if os.path.exists(local_path):
        logger.info(f"IMAGE RETRIEVAL: Cache hit for {filename} -> serving {local_path} (jpg)")
        return FileResponse(local_path, media_type="image/jpeg")
    elif os.path.exists(local_path_png):
        logger.info(f"IMAGE RETRIEVAL: Cache hit for {filename} -> serving {local_path_png} (png)")
        return FileResponse(local_path_png, media_type="image/png")

    # 2. If cache miss, download/generate the image sequentially using a global Lock
    recovered_content = None
    mime_type = "image/jpeg"
    
    # Acquire serialization lock to prevent 402 Queue full errors from concurrent Pollinations requests
    async with image_lock:
        # Re-check cache inside the lock in case another request downloaded it while we waited
        if os.path.exists(local_path):
            return FileResponse(local_path, media_type="image/jpeg")
        elif os.path.exists(local_path_png):
            return FileResponse(local_path_png, media_type="image/png")
            
        if legacy_url:
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    logger.info(f"IMAGE RECOVERY: Fetching legacy URL (Attempt {attempt+1}/{max_retries}): {legacy_url}")
                    async with httpx.AsyncClient(follow_redirects=True, timeout=20.0) as client:
                        resp = await client.get(legacy_url)
                        if resp.status_code == 200 and len(resp.content) > 5000:
                            recovered_content = resp.content
                            mime_type = resp.headers.get("content-type", "image/jpeg")
                            logger.info(f"IMAGE RECOVERY: Successfully recovered legacy image content.")
                            break
                        elif resp.status_code == 402:
                            logger.warning(f"IMAGE RECOVERY: Received 402 (Queue full) from Pollinations on attempt {attempt+1}. Retrying in 2.0s...")
                            await asyncio.sleep(2.0)
                        else:
                            logger.warning(f"IMAGE RECOVERY: Unexpected status code {resp.status_code} on attempt {attempt+1}.")
                except Exception as le:
                    logger.error(f"IMAGE RECOVERY: Failed to recover legacy URL on attempt {attempt+1}: {le}")
                    await asyncio.sleep(2.0)
                    
        if not recovered_content:
            logger.info(f"IMAGE RECOVERY: Cache miss for {filename} and legacy URL failed. Generating dynamically...")
            # Get metadata/prompt
            meta = await db_manager.get_image_metadata()
            prompt = meta.get(filename)
            if not prompt:
                prompt = "Professional clean high-fidelity editorial visual, corporate research intelligence concept"
                
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    prompt_encoded = urllib.parse.quote(prompt)
                    pollinations_url = f"https://image.pollinations.ai/prompt/{prompt_encoded}?width=1024&height=768&nologo=true&seed={random.randint(1, 100000)}"
                    logger.info(f"IMAGE RECOVERY: Generating via Pollinations AI (Attempt {attempt+1}/{max_retries}): {pollinations_url}")
                    async with httpx.AsyncClient(follow_redirects=True, timeout=20.0) as client:
                        resp = await client.get(pollinations_url)
                        if resp.status_code == 200 and len(resp.content) > 5000:
                            recovered_content = resp.content
                            mime_type = resp.headers.get("content-type", "image/jpeg")
                            logger.info(f"IMAGE RECOVERY: Successfully generated recovered image.")
                            break
                        elif resp.status_code == 402:
                            logger.warning(f"IMAGE RECOVERY: Received 402 (Queue full) from Pollinations on attempt {attempt+1}. Retrying in 2.0s...")
                            await asyncio.sleep(2.0)
                except Exception as e:
                    logger.error(f"IMAGE RECOVERY: Failed to generate on attempt {attempt+1}: {e}")
                    await asyncio.sleep(2.0)

    if recovered_content:
        # Save locally to static image cache
        try:
            os.makedirs("static/images", exist_ok=True)
            ext = "png" if "png" in mime_type.lower() else "jpg"
            save_path = os.path.join("static", "images", f"{file_hash}.{ext}")
            with open(save_path, "wb") as f:
                f.write(recovered_content)
            logger.info(f"IMAGE RETRIEVAL: Successfully saved file to local cache: {save_path}")
            return FileResponse(save_path, media_type=mime_type)
        except Exception as e:
            logger.error(f"IMAGE RETRIEVAL: Failed to save file locally: {e}")
            return Response(content=recovered_content, media_type=mime_type)
    else:
        # Generate and serve our custom high-fidelity SVG placeholder dynamically
        # to satisfy the browser retry loop and keep a beautiful, styled UI
        
        # Get dynamic description from filename (which is legacy_{encoded_url})
        try:
            if "legacy_" in filename:
                desc = filename.split("legacy_")[-1].replace(".jpg", "").replace(".png", "")
                desc = urllib.parse.unquote(desc)
                if "Topic: " in desc:
                    desc = desc.split("Topic: ")[-1]
            else:
                desc = filename
        except Exception:
            desc = "Visual Analytics"
            
        svg_content = f"""<svg width="1024" height="768" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <linearGradient id="grad" x1="0%" y1="0%" x2="100%" y2="100%">
      <stop offset="0%" style="stop-color:#2a1b18;stop-opacity:1" />
      <stop offset="100%" style="stop-color:#1c110f;stop-opacity:1" />
    </linearGradient>
  </defs>
  <rect width="1024" height="768" fill="url(#grad)" />
  <rect width="1024" height="768" fill="none" stroke="#b45309" stroke-width="4" opacity="0.2" />
  <circle cx="512" cy="320" r="70" fill="#b45309" opacity="0.1" />
  
  <!-- Sleek visual camera/graph icon -->
  <rect x="477" y="295" width="70" height="50" rx="8" fill="none" stroke="#b45309" stroke-width="3" opacity="0.5" />
  <circle cx="512" cy="320" r="14" fill="none" stroke="#b45309" stroke-width="3" opacity="0.5" />
  <path d="M500 295 L505 285 L519 285 L524 295 Z" fill="#b45309" opacity="0.4" />
  
  <text x="512" y="440" font-family="'Segoe UI', Roboto, sans-serif" font-size="20" font-weight="900" letter-spacing="4" fill="#d97706" text-anchor="middle">
    ORCHESTRA AI VISUAL
  </text>
  <text x="512" y="485" font-family="'Segoe UI', Roboto, sans-serif" font-size="15" fill="#fef3c7" text-anchor="middle" font-weight="600">
    {desc[:85] + ('...' if len(desc) > 85 else '')}
  </text>
  <text x="512" y="525" font-family="'Segoe UI', Roboto, sans-serif" font-size="12" fill="#d97706" text-anchor="middle" font-style="italic" opacity="0.7">
    Queueing live generation • High-Fidelity Asset Cache
  </text>
</svg>"""
        return Response(content=svg_content, media_type="image/svg+xml")
            
    raise HTTPException(status_code=404, detail="Image not found and recovery failed.")

# ---------------------------------------------------------------------------
# API Rate Limiter (Middleware)
# ---------------------------------------------------------------------------
@app.middleware("http")
async def rate_limit_middleware(request: Request, call_next):
    # Exclude static files and root checks to prevent locking out assets
    path = request.url.path
    if path.startswith("/static") or path == "/" or path == "/test_route":
        return await call_next(request)
        
    client_ip = request.client.host if request.client else "127.0.0.1"
    
    # Map endpoints to their corresponding sliding window limits (Requests, Window in Seconds)
    limits = {
        "/research": (5, 60),
        "/stt": (5, 60),
        "/upload": (5, 60),
        "/export/pdf": (10, 60),
        "/export/docx": (10, 60),
        "/history": (30, 60),
        "/history/upload": (30, 60),
        "/history/bulk_upload": (30, 60),
        "/history/delete": (30, 60),
        "/history/rename": (30, 60),
    }
    
    # Fallback default rate limit for any unmapped API endpoint
    limit, window = limits.get(path, (60, 60))
    
    key = f"{client_ip}:{path}"
    allowed, retry_after = await rate_limiter.is_allowed(key, limit, window)
    
    if not allowed:
        logger.warning(f"RATE LIMIT TRIGGERED: IP={client_ip} on Path={path}. Wait={retry_after}s")
        from fastapi.responses import JSONResponse
        return JSONResponse(
            status_code=429,
            headers={"Retry-After": str(retry_after)},
            content={
                "status": "Error",
                "message": f"Too many requests to {path}. Please try again in {retry_after}s.",
                "retry_after": retry_after
            }
        )
        
    return await call_next(request)


class ResearchRequest(BaseModel):
    topic: str | None = None
    query: str | None = None
    generate_images: bool = True

class DeleteReportRequest(BaseModel):
    user_id: str
    report_id: str

class RenameReportRequest(BaseModel):
    user_id: str
    report_id: str
    title: str

class PDFExportRequest(BaseModel):
    report_markdown: str
    topic: str

@app.post("/export/pdf")
async def export_pdf(req: PDFExportRequest):
    try:
        from pdf_export import generate_pdf
        pdf_bytes = await generate_pdf(req.report_markdown, req.topic)
        from fastapi.responses import Response
        return Response(content=pdf_bytes, media_type="application/pdf")
    except Exception as e:
        logger.error(f"PDF Export Error: {e}")
        return JSONResponse(status_code=500, content={"status": "Error", "message": str(e)})

@app.get("/history")
async def get_history(user_id: str, request: Request):
    from database import db_manager
    try:
        auth_header = request.headers.get("Authorization")
        token = auth_header.split(" ")[1] if auth_header else None
        history = await db_manager.get_user_history(user_id, token=token)
        return history
    except Exception as e:
        logger.error(f"History Fetch Error: {e}")
        return JSONResponse(status_code=500, content={"status": "Error", "message": str(e)})


@app.get("/")
def read_root():
    return {"status": "online", "message": "Orchestra AI Backend is Running!"}

@app.post("/stt")
async def speech_to_text(file: UploadFile = File(...)):
    """Transcribes audio using Groq Whisper-v3."""
    try:
        temp_file = f"temp_{file.filename}"
        with open(temp_file, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        with open(temp_file, "rb") as audio_file:
            translation = groq_client.audio.transcriptions.create(
                file=(temp_file, audio_file.read()),
                model="whisper-large-v3",
                response_format="json",
            )
        
        os.remove(temp_file)
        return {"text": translation.text}
    except Exception as e:
        logger.error(f"STT Error: {e}")
        return {"status": "Error", "message": str(e)}

from langchain_core.messages import HumanMessage, AIMessage

# In-memory session store
sessions = {}

@app.post("/upload")
async def upload_document(session_id: str, file: UploadFile = File(...)):
    """Handles text document uploads, generates embeddings, and indexes them in Supabase."""
    from database import db_manager
    import google.generativeai as genai
    
    file_name = file.filename or ""
    if not file_name.lower().endswith(('.txt', '.md', '.py', '.js', '.ts', '.pdf')):
        return JSONResponse(status_code=400, content={"status": "Error", "message": f"File type {file_name.split('.')[-1]} not yet supported. Please use .txt, .md, or .pdf"})

    try:
        content = await file.read()
        text = ""
        
        if file_name.lower().endswith('.pdf'):
            from pypdf import PdfReader
            import io
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        else:
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin-1")
            
        # Simple Chunking (800 chars with overlap)
        chunks = []
        chunk_size = 1000
        overlap = 200
        for i in range(0, len(text), chunk_size - overlap):
            chunk = text[i:i + chunk_size].strip()
            if len(chunk) > 20: # Ignore tiny fragments
                chunks.append(chunk)

        if not chunks:
            return {"status": "Error", "message": "Document appears to be empty or unreadable."}

        # Initialize Gemini for Embeddings
        genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
        
        for i, chunk in enumerate(chunks):
            result = genai.embed_content(
                model="models/gemini-embedding-001",
                content=chunk,
                task_type="retrieval_document",
                title=f"Doc_{file.filename}_{i}"
            )
            embedding = result['embedding']
            
            await db_manager.save_document_chunk(
                content=chunk,
                embedding=embedding,
                metadata={"source": file_name, "session_id": session_id, "chunk_index": i}
            )
            
        logger.info(f"INDEXED: {len(chunks)} chunks from {file_name}")
        return {"status": "Success", "message": f"Successfully indexed {len(chunks)} contextual fragments from {file_name}."}
    except Exception as e:
        logger.error(f"Upload/Index Error: {e}")
        return JSONResponse(status_code=500, content={"status": "Error", "message": str(e)})

@app.get("/history")
async def get_history(user_id: str, request: Request):
    """Retrieves the research history for a specific user using their session token."""
    from database import db_manager
    auth_header = request.headers.get("Authorization")
    history = await db_manager.get_user_history(user_id, auth_header)
    return history

@app.get("/test_route")
async def test_route():
    return {"status": "ok"}

@app.post("/history/upload")
async def upload_history_item(request: Request):
    """Backfills a single history item to the cloud."""
    from database import db_manager
    try:
        data = await request.json()
        auth_header = request.headers.get("Authorization")
        
        result = await db_manager.save_report(
            user_id=data.get("user_id"),
            topic=data.get("topic"),
            plan=data.get("plan", ""),
            report=data.get("report", ""),
            citations=data.get("citations", []),
            token=auth_header
        )
        return {"status": "success", "result": result}
    except Exception as e:
        logger.error(f"Backfill Error: {e}")
        return {"status": "error", "message": str(e)}

@app.post("/history/bulk_upload")
async def bulk_upload_history(request: Request):
    """Backfills multiple history items in a single request to reduce network overhead."""
    from database import db_manager
    try:
        data = await request.json()
        items = data.get("items", [])
        auth_header = request.headers.get("Authorization")
        
        results = []
        for item in items:
            res = await db_manager.save_report(
                user_id=item.get("user_id"),
                topic=item.get("topic"),
                plan=item.get("plan", ""),
                report=item.get("report", ""),
                citations=item.get("citations", []),
                token=auth_header
            )
            results.append(res)
            
        return {"status": "success", "processed": len(results)}
    except Exception as e:
        logger.error(f"Bulk Backfill Error: {e}")
        return {"status": "error", "message": str(e)}

@app.delete("/history/delete")
async def delete_history_item(request_data: DeleteReportRequest, request: Request):
    """Deletes a research report from history."""
    from database import db_manager
    try:
        auth_header = request.headers.get("Authorization")
        result = await db_manager.delete_report(
            user_id=request_data.user_id,
            report_id=request_data.report_id,
            token=auth_header
        )
        if result:
            return {"status": "success"}
        raise HTTPException(status_code=500, detail="Failed to delete report")
    except Exception as e:
        logger.error(f"Delete Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/history/rename")
async def rename_history_item(request_data: RenameReportRequest, request: Request):
    """Renames a research report in history."""
    from database import db_manager
    try:
        auth_header = request.headers.get("Authorization")
        result = await db_manager.rename_report(
            user_id=request_data.user_id,
            report_id=request_data.report_id,
            new_title=request_data.title,
            token=auth_header
        )
        if result:
            return {"status": "success"}
        raise HTTPException(status_code=500, detail="Failed to rename report")
    except Exception as e:
        logger.error(f"Rename Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.websocket("/ws/research")
async def websocket_research(websocket: WebSocket):
    """WebSocket endpoint for real-time research streaming with session support."""
    client_ip = websocket.client.host if websocket.client else "127.0.0.1"
    key = f"{client_ip}:/ws/research"
    
    allowed, retry_after = await rate_limiter.is_allowed(key, 5, 60)
    if not allowed:
        logger.warning(f"RATE LIMIT TRIGGERED (WS): IP={client_ip}. Wait={retry_after}s")
        await websocket.accept()
        await websocket.send_json({
            "type": "error",
            "message": f"Too many research queries. Please wait {retry_after} seconds before requesting again."
        })
        await websocket.close(code=1008)
        return

    await websocket.accept()
    session_id = "default"
    try:
        data = await websocket.receive_text()
        request_data = json.loads(data)
        topic = request_data.get("topic") or request_data.get("query")
        session_id = request_data.get("session_id", "default")
        depth = int(request_data.get("depth", 2))
        auth_token = request_data.get("token")
        generate_images = request_data.get("generate_images", True)
        
        if not topic:
            await websocket.send_json({"type": "error", "message": "No topic provided"})
            await websocket.close()
            return

        logger.info(f"WS START: Research on topic: {topic} (Session: {session_id}, Depth: {depth})")
        
        user_id = request_data.get("user_id", "anonymous")
        
        # Initialize or retrieve session state
        if session_id not in sessions:
            sessions[session_id] = {
                "task": topic,
                "user_id": user_id,
                "research_notes": [],
                "citations": [],
                "local_context": [],
                "revision_count": 0,
                "max_revisions": depth,
                "messages": [],
                "plan": "",
                "critic_feedback": "",
                "final_output": "",
                "generate_images": generate_images
            }
        else:
            # It's a follow-up question
            sessions[session_id]["task"] = topic
            sessions[session_id]["messages"].append(HumanMessage(content=topic))
            sessions[session_id]["max_revisions"] = depth
            # Reset notes and count for the new research cycle
            sessions[session_id]["research_notes"] = []
            sessions[session_id]["citations"] = []
            sessions[session_id]["revision_count"] = 0
            sessions[session_id]["generate_images"] = generate_images

        # Stream events from LangGraph
        async for event in graph_app.astream_events(sessions[session_id], version="v2"):
            kind = event["event"]
            
            if kind in ["on_node_start", "on_chain_start"]:
                node_name = event["name"]
                agent_map = {
                    "architect": "architect", 
                    "researcher": "researcher", 
                    "evidence": "evidence",
                    "critic": "critic", 
                    "synthesizer": "synthesizer",
                    "visualizer": "visualizer"
                }
                if node_name in agent_map:
                    await websocket.send_json({
                        "type": "node_start", 
                        "agent": agent_map[node_name]
                    })
            
            elif kind in ["on_node_end", "on_chain_end"]:
                node_name = event["name"]
                if node_name in ["architect", "researcher", "evidence", "visualizer", "critic", "synthesizer"]:
                    output = event["data"].get("output", {})
                    if not output: continue

                    # Update local session state
                    sessions[session_id].update(output)
                    
                    # Stream intermediate thoughts to frontend for Neural Matrix terminal
                    node_text = ""
                    if "messages" in output and output["messages"]:
                        node_text = getattr(output["messages"][-1], "content", str(output["messages"][-1]))
                    elif "plan" in output and output["plan"]:
                        node_text = output["plan"]
                    elif "research_notes" in output and output["research_notes"]:
                        node_text = str(output["research_notes"][-1])
                    elif "critic_feedback" in output and output["critic_feedback"]:
                        node_text = str(output["critic_feedback"])

                    if node_text:
                        # Truncate to avoid huge payloads
                        node_text = node_text[:800] + ("..." if len(node_text) > 800 else "")
                        await websocket.send_json({
                            "type": "thought",
                            "agent": node_name,
                            "content": node_text
                        })

        # Final Data Capture
        final_output = sessions[session_id].get("final_output", "").replace("APPROVED", "").strip()
        citations = sessions[session_id].get("citations", [])
        
        # Cloud Persistence (Optional based on Auth)
        user_id = sessions[session_id].get("user_id", "anonymous")
        logger.info(f"PERSISTENCE ATTEMPT: User={user_id}, Topic={topic}")
        
        if user_id != "anonymous":
            from database import db_manager
            save_result = await db_manager.save_report(
                user_id=user_id,
                topic=topic,
                plan=sessions[session_id].get("plan", ""),
                report=final_output,
                citations=citations,
                token=auth_token
            )
            logger.info(f"SAVE RESULT: {save_result}")

        # Final Report Completion - Send AFTER saving to avoid race condition
        await websocket.send_json({
            "type": "complete",
            "final_output": final_output,
            "citations": citations
        })
        
    except WebSocketDisconnect:
        logger.info(f"WebSocket client disconnected (Session: {session_id})")
    except Exception as e:
        logger.error(f"WS Error: {e}")
        try:
            await websocket.send_json({"type": "error", "message": str(e)})
        except:
            pass
    finally:
        try:
            await websocket.close()
        except:
            pass

@app.post("/research")
async def run_research(request: ResearchRequest):
    topic = request.topic or request.query
    if not topic:
        return {"status": "Error", "message": "No topic or query provided"}
    
    start_time = time.time()
    logger.info(f"START: Research on topic: {topic}")
    
    try:
        initial_state = {
            "task": topic,
            "research_notes": [],
            "revision_count": 0,
            "plan": "",
            "critic_feedback": "",
            "final_output": "",
            "generate_images": request.generate_images
        }
        
        final_state = await graph_app.ainvoke(initial_state)
        
        duration = time.time() - start_time
        logger.info(f"SUCCESS: Research completed in {duration:.2f}s for topic: {topic}")
        
        return {
            "topic": topic,
            "plan": final_state.get("plan"),
            "research_steps": final_state.get("research_notes"),
            "iterations": final_state.get("revision_count"),
            "final_output": final_state.get("final_output"),
            "status": "Success",
            "metadata": {
                "duration": round(duration, 2),
                "timestamp": time.time()
            }
        }
    except Exception as e:
        logger.error(f"FAILURE: Research failed for topic: {topic}. Error: {e}")
        return {
            "topic": topic,
            "status": "Error",
            "message": str(e)
        }

@app.post("/export/docx")
async def export_docx(request: Request):
    """Generates a systematic, well-documented DOCX file from markdown research data."""
    try:
        data = await request.json()
        topic = data.get("topic", "Research_Report")
        markdown_text = data.get("content", "")
        
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn
        import io
        import re
        import httpx
        import datetime
        import os

        def set_cell_background(cell, fill_hex):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
            tcPr.append(shd)

        def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
                node = OxmlElement(m)
                node.set(qn('w:w'), str(val))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

        def add_formatted_text(paragraph, text):
            pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))')
            parts = pattern.split(text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = paragraph.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = paragraph.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(190, 24, 74)
                elif part.startswith('[') and '](' in part and part.endswith(')'):
                    link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                    if link_match:
                        link_text = link_match.group(1)
                        run = paragraph.add_run(link_text)
                        run.font.color.rgb = RGBColor(2, 132, 199)
                        run.underline = True
                else:
                    paragraph.add_run(part)

        async def add_docx_image(doc, img_alt, img_url):
            import hashlib
            local_path = None
            
            # Determine the safe filename and check the local static cache first
            try:
                # Standardize filename format used by cache
                if "pollinations.ai" in img_url:
                    url_path = img_url.split('/')[-1].split('?')[0]
                    filename = f"legacy_{url_path}.jpg"
                else:
                    filename = os.path.basename(img_url)
                
                file_hash = hashlib.sha256(filename.encode('utf-8')).hexdigest()
                path_jpg = os.path.join("static", "images", f"{file_hash}.jpg")
                path_png = os.path.join("static", "images", f"{file_hash}.png")
                
                if os.path.exists(path_jpg):
                    local_path = path_jpg
                elif os.path.exists(path_png):
                    local_path = path_png
            except Exception as ex:
                logger.warning(f"DOCX: Failed to check local hashed cache: {ex}")
                
            if local_path:
                try:
                    logger.info(f"DOCX: Cache hit for image -> inserting {local_path} directly from disk.")
                    with open(local_path, 'rb') as lf:
                        image_data = io.BytesIO(lf.read())
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(6)
                    p_img.paragraph_format.space_after = Pt(4)
                    doc.add_picture(image_data, width=Inches(5.5))
                    
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(6)
                    r_cap = p_cap.add_run(f"Figure: {img_alt} (Source: {img_url})")
                    r_cap.font.name = 'Arial'
                    r_cap.font.size = Pt(8.5)
                    r_cap.font.italic = True
                    r_cap.font.color.rgb = RGBColor(100, 116, 139)
                    return True
                except Exception as e:
                    logger.error(f"DOCX: Failed to insert local image {local_path}: {e}")

            # Fallback to network download if not cached
            download_url = img_url
            if img_url.startswith('/static/') or img_url.startswith('static/'):
                from database import db_manager
                download_url = f"{db_manager.url}/storage/v1/object/public/images/{filename}"
            
            if download_url.startswith('http://') or download_url.startswith('https://'):
                try:
                    async with httpx.AsyncClient() as client:
                        resp = await client.get(download_url, timeout=10.0)
                        if resp.status_code == 200:
                            image_data = io.BytesIO(resp.content)
                            p_img = doc.add_paragraph()
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_img.paragraph_format.space_before = Pt(6)
                            p_img.paragraph_format.space_after = Pt(4)
                            doc.add_picture(image_data, width=Inches(5.5))
                            
                            p_cap = doc.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_cap.paragraph_format.space_after = Pt(6)
                            r_cap = p_cap.add_run(f"Figure: {img_alt} (Source: {img_url})")
                            r_cap.font.name = 'Arial'
                            r_cap.font.size = Pt(8.5)
                            r_cap.font.italic = True
                            r_cap.font.color.rgb = RGBColor(100, 116, 139)
                            return True
                except Exception as e:
                    logger.warning(f"DOCX: Could not download image {download_url}: {e}")
                    
            p_placeholder = doc.add_paragraph()
            p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_placeholder.paragraph_format.space_before = Pt(6)
            p_placeholder.paragraph_format.space_after = Pt(4)
            run_pl = p_placeholder.add_run(f"[IMAGE NOT AVAILABLE: {img_alt}]")
            run_pl.font.name = 'Arial'
            run_pl.font.size = Pt(9.5)
            run_pl.font.color.rgb = RGBColor(190, 24, 74)
            return False

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        normal_style = doc.styles['Normal']
        font = normal_style.font # type: ignore
        font.name = 'Arial'
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(30, 41, 59)
        normal_style.paragraph_format.line_spacing = 1.15 # type: ignore
        normal_style.paragraph_format.space_after = Pt(6) # type: ignore
        
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "Orchestra AI Intelligence Report  |  " + topic
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.name = 'Arial'
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(100, 116, 139)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Prepared by Orchestra AI agent swarm  |  Confidential"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.runs[0].font.name = 'Arial'
        fp.runs[0].font.size = Pt(8.5)
        fp.runs[0].font.color.rgb = RGBColor(148, 163, 184)
        
        p_top = doc.add_paragraph()
        p_top.paragraph_format.space_before = Pt(120)
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sub = p_sub.add_run("ORCHESTRA AI EXECUTIVE REPORT")
        r_sub.font.name = 'Arial'
        r_sub.font.size = Pt(10)
        r_sub.font.bold = True
        r_sub.font.color.rgb = RGBColor(8, 145, 178)
        p_sub.paragraph_format.space_after = Pt(12)
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_title = p_title.add_run(topic)
        r_title.font.name = 'Arial'
        r_title.font.size = Pt(28)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(30, 41, 59)
        p_title.paragraph_format.space_after = Pt(24)
        
        div_table = doc.add_table(rows=1, cols=1)
        div_table.alignment = 1 # type: ignore
        div_cell = div_table.cell(0, 0)
        div_cell.width = Inches(2.5)
        set_cell_background(div_cell, "0891B2")
        set_cell_margins(div_cell, top=20, bottom=20, left=0, right=0)
        for p in div_cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.space_before = Pt(180)
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_meta = p_meta.add_run(
            f"DATE OF GENERATION: {datetime.date.today().strftime('%B %d, %Y')}\n"
            f"PREPARED FOR: Executive Review\n"
            f"PLATFORM: Orchestra AI Swarm Analytics"
        )
        run_meta.font.name = 'Arial'
        run_meta.font.size = Pt(9)
        run_meta.font.color.rgb = RGBColor(100, 116, 139)
        p_meta.paragraph_format.space_after = Pt(0)
        
        doc.add_page_break()
        
        lines = markdown_text.split('\n')
        i = 0
        in_code_block = False
        code_content = []
        in_table = False
        table_lines = []
        
        while i < len(lines):
            line = lines[i]
            stripped_line = line.strip()
            
            if stripped_line.startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_content = []
                else:
                    in_code_block = False
                    code_text = '\n'.join(code_content)
                    c_table = doc.add_table(rows=1, cols=1)
                    c_cell = c_table.cell(0, 0)
                    set_cell_background(c_cell, "F8FAFC")
                    set_cell_margins(c_cell, top=120, bottom=120, left=180, right=180)
                    tcPr = c_cell._tc.get_or_add_tcPr()
                    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="0891B2"/><w:top w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
                    tcPr.append(borders)
                    cp = c_cell.paragraphs[0]
                    cp.paragraph_format.space_before = Pt(0)
                    cp.paragraph_format.space_after = Pt(0)
                    run_c = cp.add_run(code_text)
                    run_c.font.name = 'Consolas'
                    run_c.font.size = Pt(9.0)
                    run_c.font.color.rgb = RGBColor(71, 85, 105)
                i += 1
                continue
            if in_code_block:
                code_content.append(line)
                i += 1
                continue
            if stripped_line.startswith('|'):
                in_table = True
                table_lines.append(line)
                i += 1
                continue
            elif in_table:
                in_table = False
                parsed_rows = []
                for tl in table_lines:
                    t_stripped = tl.strip()
                    if t_stripped.startswith('|'): t_stripped = t_stripped[1:]
                    if t_stripped.endswith('|'): t_stripped = t_stripped[:-1]
                    cols = [c.strip() for c in t_stripped.split('|')]
                    parsed_rows.append(cols)
                filtered_rows = [r for r in parsed_rows if not all(re.match(r'^[\s\-:]+$', cell) for cell in r)]
                if filtered_rows:
                    num_cols = max(len(r) for r in filtered_rows)
                    table_word = doc.add_table(rows=len(filtered_rows), cols=num_cols)
                    table_word.alignment = 1 # type: ignore
                    for r_idx, row_data in enumerate(filtered_rows):
                        for c_idx, cell_value in enumerate(row_data):
                            w_cell = table_word.cell(r_idx, c_idx)
                            set_cell_margins(w_cell, top=100, bottom=100, left=120, right=120)
                            tcPr = w_cell._tc.get_or_add_tcPr()
                            borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/><w:left w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/><w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/></w:tcBorders>')
                            tcPr.append(borders)
                            p_cell = w_cell.paragraphs[0]
                            if r_idx == 0:
                                set_cell_background(w_cell, "0F172A")
                                add_formatted_text(p_cell, cell_value)
                                for run in p_cell.runs: run.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                if r_idx % 2 == 1: set_cell_background(w_cell, "F8FAFC")
                                add_formatted_text(p_cell, cell_value)
                table_lines = []
                continue
            if stripped_line.startswith('# '):
                h = doc.add_heading(level=1)
                run_h = h.add_run(stripped_line[2:])
                run_h.font.name = 'Arial'; run_h.font.size = Pt(16); run_h.bold = True; run_h.font.color.rgb = RGBColor(8, 145, 178)
            elif stripped_line.startswith('## '):
                h = doc.add_heading(level=2)
                run_h = h.add_run(stripped_line[3:])
                run_h.font.name = 'Arial'; run_h.font.size = Pt(13); run_h.bold = True; run_h.font.color.rgb = RGBColor(71, 85, 105)
            elif stripped_line.startswith('### '):
                h = doc.add_heading(level=3)
                run_h = h.add_run(stripped_line[4:])
                run_h.font.name = 'Arial'; run_h.font.size = Pt(11); run_h.bold = True; run_h.font.color.rgb = RGBColor(30, 41, 59)
            elif stripped_line.startswith('- ') or stripped_line.startswith('* ') or re.match(r'^\d+\. ', stripped_line):
                style = 'List Bullet' if not re.match(r'^\d+\. ', stripped_line) else 'List Number'
                p = doc.add_paragraph(style=style)
                add_formatted_text(p, re.sub(r'^\d+\. ', '', stripped_line[2:]) if 'List Number' in style else stripped_line[2:])
            elif '![' in stripped_line and '](' in stripped_line:
                img_match = re.search(r'!\[(.*?)\]\((.*?)\)', stripped_line)
                if img_match: await add_docx_image(doc, img_match.group(1), img_match.group(2))
            elif stripped_line:
                p = doc.add_paragraph(); add_formatted_text(p, line)
            i += 1
        target_stream = io.BytesIO()
        doc.save(target_stream); target_stream.seek(0)
        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            target_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Orchestra_Research_{topic.replace(' ', '_')}.docx"}
        )
    except Exception as e:
        logger.error(f"DOCX Export Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=8000)

